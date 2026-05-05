"""
Собирает презентационное письмо ООО «Деловые Решения» в .docx
по образцу письма ООО «АвтоДорСтрой».
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


FONT = "Times New Roman"
OUT_PATH = Path(__file__).resolve().parent.parent / "Презентационное письмо ООО Деловые Решения.docx"


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def add_paragraph(doc, text="", *, size=12, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0, space_before=0,
                  first_line_indent=None, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_runs(p, parts):
    """parts: list of (text, {bold,italic,size}) tuples."""
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            size=opts.get("size", 12),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def write_cell(cell, text, *, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
               italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    p.text = ""
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    set_cell_borders(cell)


def add_field(paragraph, instr_text):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, size=10)
    run._element.append(fld_char1)
    run._element.append(instr)
    run._element.append(fld_char2)


def build_header(section):
    header = section.header
    header.is_linked_to_previous = False
    while header.paragraphs and len(header.paragraphs) > 1:
        p = header.paragraphs[-1]
        p._element.getparent().remove(p._element)
    p0 = header.paragraphs[0]
    p0.text = ""
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(0)
    p0.paragraph_format.line_spacing = 1.0
    run = p0.add_run("ООО «Деловые Решения»")
    set_run_font(run, size=11, bold=True)

    lines = [
        "ИНН 7731371885 / КПП 773101001  •  ОГРН 1177746584817",
        "Юр. адрес: 121596, г. Москва, вн.тер.г. Муниципальный округ Можайский, ул. Горбунова, д. 2, стр. 3, пом. 189",
        "АО «АЛЬФА-БАНК»  •  Р/с 40702810001300022656  •  К/с 30101810200000000593  •  БИК 044525593",
        "Тел.: +7 (495) 210-85-85  •  e-mail: info-deloresh@mail.ru  •  сайт: deloresh.ru",
    ]
    for line in lines:
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, size=9)

    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.text = ""
    run = p.add_run("стр. ")
    set_run_font(run, size=9)
    add_field(p, "PAGE")
    run2 = p.add_run(" из ")
    set_run_font(run2, size=9)
    add_field(p, "NUMPAGES")


# ---------------------------------------------------------------------------
# Контент письма
# ---------------------------------------------------------------------------

REFERENCES = [
    (
        "ООО «ДСК ПАРТНЕР»",
        [
            (
                "Выполнение работ по благоустройству территорий Юго-Восточного "
                "административного округа города Москвы и расположенных на них "
                "объектах по Договору субподряда №135/СУБ/ДР-24 от 09.04.2024 г. "
                "на объекте по адресу: Маршала Голованова улица (на всём протяжении).",
                "159 413 188,00 руб.",
            ),
        ],
    ),
    (
        "ООО «ПК ПАРТНЕР»",
        [
            (
                "Выполнение работ по благоустройству территорий административных "
                "округов города Москвы и расположенных на них объектах (СВАО) по "
                "Договору № АНО2/СУБ-01/24 от 23.09.2024 г. на объекте по адресу: "
                "Проектируемый проезд №226 (на всём протяжении).",
                "63 181 847,02 руб.",
            ),
        ],
    ),
    (
        "ООО «ПСК СТРОЙМОНОЛИТ»",
        [
            (
                "Выполнение работ по благоустройству территорий по адресам: "
                "г. Мариуполь, пр. Металлургов, д. 215, д. 213, д. 237; "
                "пр. Ильича, д. 140. Согласно договору подряда №2023-03/22(П-Мар) "
                "от 22.03.2023 г. и КС-3 №1 от 20.07.2023 г.",
                "33 844 365,50 руб.",
            ),
            (
                "Выполнение подрядных работ по благоустройству территории детского "
                "сада, расположенного по адресу: г. Москва, поселение Филимоновское, "
                "ул. Харлампиева. Согласно договору №2023/0601 (ПСК/Фил/благ/П) "
                "от 01.06.2023 г.",
                "28 574 600,00 руб.",
            ),
        ],
    ),
    (
        "ООО «РОУД ГРУПП»",
        [
            (
                "Выполнение работ по благоустройству территорий Юго-Западного "
                "административного округа города Москвы и расположенных на них "
                "объектах по Договору №СП-1174/136-2024-АД от 15.04.2024 г. на "
                "объекте: Нахимовский проспект (на всём протяжении).",
                "80 544 929,45 руб.",
            ),
            (
                "Выполнение работ на объектах дорожного хозяйства территорий "
                "Новомосковского (НАО) и Троицкого (ТАО) административных округов "
                "города Москвы по Договору №1277/СП-500ТР-25 от 03.06.2025 г. "
                "Проезд вдоль ЖК «Бунинский» (на всём протяжении).",
                "341 404 460,96 руб.",
            ),
            (
                "Выполнение работ на объектах дорожного хозяйства территорий "
                "Западного, Новомосковского и Южного административных округов "
                "города Москвы по Договору №1295/СП-528/295/ВР-25 от 03.06.2025 г. "
                "Медынская улица (на всём протяжении).",
                "117 165 398,00 руб.",
            ),
            (
                "Выполнение комплекса работ по благоустройству территории на "
                "объекте: «Реконструкция подъезда к парку Малевича в Одинцовском "
                "городском округе Московской области». Согласно Договору "
                "субподряда №1122-1/2022-66/СП от 17.07.2023 г.",
                "45 259 812,00 руб.",
            ),
        ],
    ),
    (
        "АО «РЕМОНТЁР»",
        [
            (
                "Выполнение работ по улучшению пешеходной и транспортной "
                "доступности города Москвы (Лот 26) согласно Договору субподряда "
                "№СП-294Л5-О2-4-ГКУДКРП4 от 12.02.2024 г. на объектах по адресу: "
                "г. Москва, ул. А. Монаховой.",
                "363 683 641,00 руб.",
            ),
            (
                "Выполнение работ по улучшению пешеходной и транспортной "
                "доступности города Москвы (Лот 64) согласно Договору субподряда "
                "№СП-384/15-02-26-ГКУДКР/24 от 23.08.2024 г. на объектах по адресу: "
                "г. Москва, ул. А. Монаховой.",
                "298 782 572,83 руб.",
            ),
            (
                "Выполнение работ на объектах дорожного хозяйства территорий "
                "Юго-Западного административного округа (ЮЗАО) города Москвы. "
                "Дороги 26 квартала ЮЗАО (на всём протяжении). Согласно Договору "
                "субподряда №СП-231/221/2023-АД от 23.05.2023 г.",
                "131 108 519,65 руб.",
            ),
            (
                "Выполнение работ на объектах дорожного хозяйства территорий "
                "Западного административного округа (ЗАО) города Москвы. Проезд "
                "за БЦ «Вернадский» (на всём протяжении). Согласно Договору "
                "субподряда №СП-236/220/2023-АД от 23.05.2023 г.",
                "108 433 210,11 руб.",
            ),
            (
                "Выполнение работ по благоустройству территории ГБУ «Миграционный "
                "центр», расположенной по адресу: г. Москва, ТиНАО, "
                "пос. Вороновское, Варшавское шоссе, 64-й км, д. 1, стр. 59. "
                "Согласно Договору субподряда №СП-244/15-02-29-ГКУДКР/23 "
                "от 16.06.2023 г.",
                "33 291 286,39 руб.",
            ),
        ],
    ),
    (
        "ООО «СТРОЙ БИЗНЕС ИНВЕСТ»",
        [
            (
                "Выполнение работ по благоустройству территорий административных "
                "округов города Москвы и расположенных на них объектах согласно "
                "Договору субподряда №ПК-СБИ-25/ДР от 03.06.2025 г. на объектах "
                "по адресу: г. Москва, ул. Элеваторная, ул. Кожуховская.",
                "255 000 000,00 руб.",
            ),
        ],
    ),
    (
        "ООО «ТРИУМФ»",
        [
            (
                "Выполнение работ по улучшению пешеходной и транспортной "
                "доступности на территориях, прилегающих к станции Московского "
                "метрополитена БКЛ «Новаторская». Согласно Договору субподряда "
                "№216/32211556928/СП от 23.08.2022 г.",
                "248 644 042,65 руб.",
            ),
        ],
    ),
    (
        "ООО «МИГДОРГРУПП»",
        [
            (
                "Выполнение работ по приведению в нормативное состояние "
                "асфальтобетонного покрытия и ремонт бортового камня на территории "
                "ДНР, г. Донецк, согласно Договору субподряда "
                "№10993/МИГ-ДЕЛРЕШ-2026 от 04.03.2026 г. на объекте по адресу: "
                "г. Донецк, ул. Краснооктябрьская, ул. Логутенко.",
                "233 000 000,00 руб.",
            ),
        ],
    ),
]

EQUIPMENT = [
    (
        "Малотоннажные автомобили",
        [
            ("ГАЗ А22R32 (грузовой бортовой)", 2016),
            ("ГАЗ-А22R32 (платформа с каркасом и тентом)", 2014),
            ("ГАЗ-А22R32 (грузовой с бортом)", 2014),
            ("Газель GAZelle NEXT", 2021),
            ("ГАЗ 2705 (грузовой фургон)", 2012),
            ("ГАЗ-А22R32 (грузовой бортовой)", 2016),
            ("Ford Transit (техпомощь)", 2019),
            ("Ford Transit 22278С (грузопассажирский)", 2016),
            ("Ford Transit 3227AR двойная кабина (бортовая платформа)", 2019),
            ("Ford Transit 3227AR двойная кабина", 2019),
        ],
    ),
    (
        "Автобусы",
        [
            ("Газель GAZelle NN", 2021),
            ("Ford Transit", 2019),
            ("Ford Transit (автобус)", 2018),
        ],
    ),
    (
        "Автомобили специальные",
        [
            ("Грузовой бортовой, оснащённый краном-манипулятором TH-T13, FOTON 3006A7Н", 2024),
            ("Комбинированная дорожная машина КО-806-01", 2021),
            ("Автогудронатор РЕГИОН 45 АС-С41R", 2021),
        ],
    ),
    (
        "Самосвалы",
        [
            ("Автосамосвал КАМАЗ К3340 6520-53", 2020),
            ("Автосамосвал КАМАЗ К3340 6520-53", 2019),
            ("Автосамосвал КАМАЗ К3340 6520-53", 2020),
            ("Автосамосвал КАМАЗ К3340 6520-53", 2020),
            ("Автосамосвал КАМАЗ К3340 6520-53", 2020),
            ("Автосамосвал КАМАЗ К3340 6520-53", 2019),
        ],
    ),
    (
        "Седельные тягачи",
        [
            ("Sitrak C7H ZZ4186V361HE", 2023),
            ("SCANIA Р400 СА6Х4НSZ (грузовой тягач седельный)", 2016),
        ],
    ),
    (
        "Полуприцепы (прицепы)",
        [
            ("Specpricep 9942L3 (специализированный полуприцеп-тяжеловоз)", 2021),
            ("Прицеп Sinanli Tanker ST4FLF", 2023),
            ("Прицеп 71491-0000010-01", 2020),
        ],
    ),
    (
        "Фронтальные погрузчики",
        [
            ("Погрузчик фронтальный ENSIGN YX635", 2022),
            ("Погрузчик фронтальный ENSIGN YX635", 2022),
            ("Фронтальный погрузчик XCMG ZL30GV", 2022),
        ],
    ),
    (
        "Мини-погрузчики",
        [
            ("Погрузчик фронтальный CASE SR220", 2018),
            ("Колёсный мини-погрузчик CASE SR200B", 2022),
        ],
    ),
    (
        "Экскаваторы-погрузчики",
        [
            ("Экскаватор-погрузчик Case 580T", 2018),
            ("Экскаватор-погрузчик Caterpillar 428", 2020),
            ("Экскаватор-погрузчик Caterpillar 432F2", 2018),
            ("МУП-351.ТСТ (машина для коммунального хозяйства)", 2020),
        ],
    ),
    (
        "Экскаваторы",
        [
            ("Гусеничный экскаватор SANY SY55C", 2022),
            ("Гусеничный экскаватор SANY SY75C", 2022),
            ("Колёсный экскаватор SANY SY155W (158)", 2022),
            ("Колёсный экскаватор SANY SY155W (368)", 2022),
            ("Экскаватор SANY SY155W", 2022),
            ("Самоходная машина SANY SY155W", 2022),
            ("Экскаватор Caterpillar Inc. 318CL", 2005),
        ],
    ),
    (
        "Катки",
        [
            ("Каток вибрационный LIUGONG CLG6614E", 2022),
            ("Каток дорожный AMMANN ARX45-2", 2022),
            ("Каток дорожный AMMANN ARX45-2", 2022),
            ("Каток дорожный HD110 HAMM AG", 2013),
            ("Каток дорожный самоходный ZDM DM-10-VD", 2024),
            ("Каток дорожный самоходный ZDM ZDM-10-VC", 2024),
            ("Каток дорожный самоходный вибрационный AMMANN ARX 26 K", 2019),
            ("СМ Каток HAMM HD 090V", 2005),
            ("СМ Каток дорожный BW 161 AD-4", 2010),
        ],
    ),
    (
        "Асфальтоукладчики",
        [
            ("Асфальтоукладчик VOGELE SUPER 1900-2 (6 м)", 2011),
            ("Асфальтоукладчик VOGELE SUPER 1900-2/1 (5 м)", 2012),
            ("Асфальтоукладчик VOGELE SUPER 1300-2", 2008),
        ],
    ),
    (
        "Фрезы",
        [
            ("Фреза дорожная WIRTGEN W210", 2012),
            ("Холодная дорожная фреза WIRTGEN W200", 2010),
        ],
    ),
]


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT)

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    build_header(section)
    build_footer(section)

    add_paragraph(
        doc, "ПРЕЗЕНТАЦИОННОЕ ПИСЬМО",
        size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=12, line_spacing=1.0,
    )

    info_pairs = [
        ("Наименование предприятия — ", "Общество с ограниченной ответственностью «Деловые Решения»."),
        ("Сокращённое наименование — ", "ООО «Деловые Решения»."),
        ("ИНН / КПП — ", "7731371885 / 773101001."),
        ("ОГРН — ", "1177746584817."),
        ("ОКВЭД / ОКПО — ", "42.11 / 15996287."),
        ("Юридический адрес: ", "121596, г. Москва, вн.тер.г. Муниципальный округ Можайский, ул. Горбунова, д. 2, стр. 3, пом. 189."),
        ("Фактический адрес: ", "совпадает с юридическим."),
        ("Адрес электронной почты: ", "info-deloresh@mail.ru."),
        ("Сайт: ", "deloresh.ru."),
        ("Контактный телефон: ", "+7 (495) 210-85-85."),
    ]
    for label, value in info_pairs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.line_spacing = 1.15
        add_runs(p, [(label, {"bold": True, "size": 12}), (value, {"size": 12})])

    add_paragraph(doc, space_after=6, line_spacing=1.0)

    about_paragraphs = [
        "Общество с ограниченной ответственностью «Деловые Решения» — предприятие, "
        "специализирующееся на комплексном благоустройстве территорий и "
        "дорожно-строительных работах в г. Москве, Московской области и других "
        "регионах Российской Федерации с применением материалов и технологий нового "
        "поколения. Мы выполняем полный комплекс работ — от подготовительного этапа "
        "и до сдачи объекта в эксплуатацию. За время своей работы зарекомендовали "
        "себя как надёжный деловой партнёр для государственных заказчиков и крупных "
        "подрядных организаций.",

        "Высококвалифицированные специалисты нашей компании, собственный парк "
        "дорожно-строительной техники (асфальтоукладочные комплексы, дорожные катки, "
        "холодные фрезы, экскаваторы и экскаваторы-погрузчики, фронтальные и "
        "мини-погрузчики, самосвальная и седельная техника, специализированные "
        "машины и средства малой механизации), высокое качество применяемых "
        "материалов и технологий — всё это даёт нам возможность гарантировать "
        "качество, безопасность и сроки выполнения работ на объектах любой сложности.",
    ]
    for text in about_paragraphs:
        add_paragraph(
            doc, text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=1.0, space_after=6, line_spacing=1.15,
        )

    add_paragraph(
        doc, "Мы выполняем следующие виды работ:",
        size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=6, space_after=4, line_spacing=1.15,
    )
    work_types = [
        "комплексное благоустройство территорий административных округов г. Москвы, "
        "Московской области и иных регионов РФ;",
        "благоустройство дворовых и общественных территорий, парков, скверов, бульваров;",
        "работы по улучшению пешеходной и транспортной доступности территорий, "
        "прилегающих к станциям Московского метрополитена и иным объектам "
        "транспортной инфраструктуры;",
        "ремонт, реконструкция и устройство автомобильных дорог общего пользования, "
        "проездов и тротуаров;",
        "работы по приведению в нормативное состояние асфальтобетонного покрытия, "
        "ямочный ремонт;",
        "устройство и ремонт бортового камня, укладка тротуарной плитки и брусчатки;",
        "земляные работы, устройство оснований дорог и тротуаров, водоотводов и дренажей;",
        "благоустройство территорий объектов социальной инфраструктуры "
        "(детские сады, миграционные центры и иные объекты);",
        "работы в качестве генерального подрядчика и субподрядчика по контрактам на "
        "территории г. Москвы, ТиНАО, Московской области, Донецкой Народной "
        "Республики и г. Мариуполя.",
    ]
    for item in work_types:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=12)

    add_paragraph(
        doc,
        "В т. ч. силами нашей организации выполнены следующие работы за период "
        "2022–2026 гг.:",
        size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=10, space_after=4, line_spacing=1.15,
    )

    for company, works in REFERENCES:
        add_paragraph(
            doc, company,
            size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=6, space_after=2, line_spacing=1.15,
        )
        for idx, (text, amount) in enumerate(works, start=1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.first_line_indent = Cm(1.0)
            prefix = f"{idx}) " if len(works) > 1 else ""
            add_runs(p, [(prefix + text, {"size": 12})])

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.line_spacing = 1.15
            p2.paragraph_format.first_line_indent = Cm(1.0)
            add_runs(
                p2,
                [
                    ("Сумма договора ", {"size": 12}),
                    (amount, {"size": 12, "bold": True}),
                ],
            )

    add_paragraph(
        doc,
        "ООО «Деловые Решения» располагает следующими ресурсами для выполнения работ:",
        size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6, line_spacing=1.15,
    )

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Cm(1.0), Cm(8.5), Cm(1.8), Cm(2.5), Cm(1.8), Cm(1.6)]
    headers = ["№", "Наименование, марка, модель", "Год выпуска",
               "Местонахождение", "Право", "Состояние"]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = col_widths[i]
        write_cell(cell, headers[i], bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "D9D9D9")

    counter = 0
    for category, items in EQUIPMENT:
        cat_row = table.add_row().cells
        merged = cat_row[0].merge(cat_row[-1])
        write_cell(
            merged, category, bold=True, italic=True, size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(merged, "F2F2F2")

        for name, year in items:
            counter += 1
            row = table.add_row().cells
            row[0].width = col_widths[0]
            row[1].width = col_widths[1]
            row[2].width = col_widths[2]
            row[3].width = col_widths[3]
            row[4].width = col_widths[4]
            row[5].width = col_widths[5]
            write_cell(row[0], str(counter), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            write_cell(row[1], name, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            write_cell(row[2], str(year), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            write_cell(row[3], "г. Москва", size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            write_cell(row[4], "в собств.", size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            write_cell(row[5], "отличное", size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, space_before=18, space_after=0, line_spacing=1.0)

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig.paragraph_format.space_before = Pt(24)
    sig.paragraph_format.space_after = Pt(0)
    sig.paragraph_format.line_spacing = 1.15
    add_runs(
        sig,
        [
            ("Генеральный директор\n", {"bold": True, "size": 12}),
            ("ООО «Деловые Решения»", {"bold": True, "size": 12}),
            ("\t\t_____________________ \t/Арамян Н. Г./", {"size": 12}),
        ],
    )

    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mp.paragraph_format.space_before = Pt(24)
    mp.paragraph_format.line_spacing = 1.15
    run = mp.add_run("М. П.")
    set_run_font(run, size=12, bold=True)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build()
    print(f"OK: {out}")
